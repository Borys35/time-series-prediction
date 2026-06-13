import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
VISUALS = ROOT / "artifacts" / "visualizations"
RESULTS = ROOT / "artifacts" / "results"
OUTPUT = ROOT / "docs" / "Sprawozdanie_TEP_MPC.docx"

NAVY = "173B63"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "5D6670"
LIGHT = "F4F6F9"
TABLE_HEADER = "E8EEF5"
WHITE = "FFFFFF"
CONTENT_DXA = 9060


def set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = margins.find(qn(f"w:{name}"))
        if element is None:
            element = OxmlElement(f"w:{name}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    properties = table._tbl.tblPr

    width = properties.find(qn("w:tblW"))
    if width is None:
        width = OxmlElement("w:tblW")
        properties.append(width)
    width.set(qn("w:w"), str(sum(widths)))
    width.set(qn("w:type"), "dxa")

    indent = properties.find(qn("w:tblInd"))
    if indent is None:
        indent = OxmlElement("w:tblInd")
        properties.append(indent)
    indent.set(qn("w:w"), "120")
    indent.set(qn("w:type"), "dxa")

    layout = properties.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        properties.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for value in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(value))
        grid.append(column)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index] / 1440)
            properties = cell._tc.get_or_add_tcPr()
            cell_width = properties.find(qn("w:tcW"))
            if cell_width is None:
                cell_width = OxmlElement("w:tcW")
                properties.append(cell_width)
            cell_width.set(qn("w:w"), str(widths[index]))
            cell_width.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run_font(run, size=None, bold=None, italic=None, color=None, name="Calibri") -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Strona ")
    set_run_font(run, size=9, color=MUTED)
    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    run._r.append(field_begin)
    run._r.append(instruction)
    run._r.append(field_end)


def add_heading(doc, text: str, level: int = 1) -> None:
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True
    paragraph.add_run(text)


def add_body(doc, text: str, bold_prefix: str | None = None) -> None:
    paragraph = doc.add_paragraph(style="Normal")
    if bold_prefix and text.startswith(bold_prefix):
        first = paragraph.add_run(bold_prefix)
        first.bold = True
        paragraph.add_run(text[len(bold_prefix):])
    else:
        paragraph.add_run(text)


def add_formula(doc, formula: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(7)
    run = paragraph.add_run(formula)
    set_run_font(run, size=11, italic=True, color=NAVY, name="Cambria Math")


def add_bullets(doc, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.add_run(item)


def add_numbered(doc, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        paragraph.add_run(item)


def add_figure(doc, filename: str, caption: str, width_cm: float = 15.4) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run()
    run.add_picture(str(VISUALS / filename), width=Cm(width_cm))

    caption_paragraph = doc.add_paragraph(style="Caption")
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph.paragraph_format.space_after = Pt(8)
    caption_paragraph.add_run(caption)


def add_table(doc, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, TABLE_HEADER)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(header)
        set_run_font(run, size=9.5, bold=True, color=NAVY)
    for row_values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_values):
            paragraph = cells[index].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if index == 0 else WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(str(value))
            set_run_font(run, size=9.3)
            set_cell_margins(cells[index])
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_table_geometry(table, widths)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)
    section.different_first_page_header_footer = True

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    normal.paragraph_format.widow_control = True

    heading_values = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in heading_values.items():
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    caption = styles["Caption"]
    caption.font.name = "Calibri"
    caption._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor.from_string(MUTED)

    for list_name in ("List Bullet", "List Number"):
        style = styles[list_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Cm(0.95)
        style.paragraph_format.first_line_indent = Cm(-0.46)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.208

    header = section.header
    header_paragraph = header.paragraphs[0]
    header_paragraph.text = "Hurtownie Danych | Prognozowanie TEP i sterowanie MPC"
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in header_paragraph.runs:
        set_run_font(run, size=9, color=MUTED)

    footer = section.footer
    add_page_field(footer.paragraphs[0])

    first_header = section.first_page_header
    first_header.paragraphs[0].text = ""
    first_footer = section.first_page_footer
    first_footer.paragraphs[0].text = ""

def add_cover(doc: Document) -> None:
    for _ in range(5):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(8)

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = kicker.add_run("SPRAWOZDANIE PROJEKTOWE")
    set_run_font(run, size=11, bold=True, color=BLUE)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(14)
    title.paragraph_format.space_after = Pt(10)
    run = title.add_run("Prognozowanie i analiza szeregów czasowych\noraz sterowanie MPC dla procesu Tennessee Eastman")
    set_run_font(run, size=25, bold=True, color=NAVY)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(55)
    run = subtitle.add_run("Projekt z przedmiotu Hurtownie Danych")
    set_run_font(run, size=14, color=MUTED)

    authors = doc.add_paragraph()
    authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
    authors.paragraph_format.space_after = Pt(5)
    run = authors.add_run("Maksymilian Lech, nr indeksu 280136")
    set_run_font(run, size=12, bold=True, color=NAVY)

    authors2 = doc.add_paragraph()
    authors2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = authors2.add_run("Borys Kaczmarek, nr indeksu 280068")
    set_run_font(run, size=12, bold=True, color=NAVY)

    for _ in range(5):
        doc.add_paragraph()

    place = doc.add_paragraph()
    place.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = place.add_run("Wrocław, 2026")
    set_run_font(run, size=11, color=MUTED)
    doc.add_page_break()


def add_front_matter(doc: Document) -> None:
    add_heading(doc, "Spis treści", 1)
    entries = [
        "1. Wprowadzenie",
        "2. Zbiór danych Tennessee Eastman Process",
        "3. Przygotowanie danych i budowa modelu",
        "4. Estymacja stanu i prognozowanie",
        "5. Ocena jakości prognoz",
        "6. Regulator MPC",
        "7. Architektura aplikacji i warstwa danych",
        "8. Pętla zamknięta",
        "9. Implementacja i testy",
        "10. Ograniczenia i możliwe rozszerzenia",
        "11. Wnioski",
        "Bibliografia",
    ]
    for entry in entries:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Cm(0.5)
        paragraph.paragraph_format.space_after = Pt(4)
        run = paragraph.add_run(entry)
        set_run_font(run, size=11, color=NAVY)
    doc.add_page_break()

    add_heading(doc, "Streszczenie", 1)
    add_body(
        doc,
        "Celem projektu było przygotowanie kompletnego, ale możliwie prostego systemu do analizy wielowymiarowych szeregów czasowych procesu Tennessee Eastman. System przetwarza kolejne próbki, estymuje ukryty stan filtrem Kalmana, prognozuje 41 pomiarów na jeden krok oraz wyznacza ograniczone sterowanie MPC. PostgreSQL przechowuje pomiary, wejścia, prognozy i wyniki, natomiast MQTT służy do przekazywania zdarzeń pomiędzy symulatorem a kontrolerem. Model oceniono na dziesięciu niezależnych przebiegach testowych. Średni NRMSE wyniósł 0,786 wobec 0,952 dla prognozy naiwnej, co oznacza poprawę o około 17,4%. Przygotowano także tryb pętli zamkniętej, w którym sterowanie MPC wpływa na kolejne stany symulowanego procesu.",
    )
    add_body(
        doc,
        "Słowa kluczowe: szeregi czasowe, Tennessee Eastman Process, PCA, filtr Kalmana, model przestrzeni stanów, MPC, PostgreSQL, MQTT.",
    )


def build_report() -> None:
    summary = json.loads((RESULTS / "podsumowanie_wynikow.json").read_text(encoding="utf-8"))
    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_front_matter(doc)

    add_heading(doc, "1. Wprowadzenie", 1)
    add_body(
        doc,
        "Analiza szeregów czasowych jest szczególnie ważna w systemach przemysłowych, ponieważ kolejne pomiary są ze sobą powiązane przez fizyczną dynamikę procesu. Temperatura, ciśnienie lub przepływ nie zmieniają się niezależnie w każdej chwili. Na ich przyszłe wartości wpływają wcześniejszy stan instalacji oraz decyzje sterujące. W projekcie wykorzystano dane Tennessee Eastman Process, które tworzą wielowymiarowy zbiór 41 pomiarów i 11 wejść sterujących.",
    )
    add_body(
        doc,
        "Pierwsza wersja projektu łączyła dane TEP, PostgreSQL, MQTT i regulator MPC, lecz model był używany w innej skali niż podczas treningu, sterowania nie miały ograniczeń, a symulator nie reagował na decyzje regulatora. Aktualna wersja zachowuje pierwotny zakres technologiczny, ale porządkuje cały przepływ i dodaje właściwą ocenę prognoz. Założeniem nie było stworzenie systemu produkcyjnego, tylko poprawnego i możliwego do wyjaśnienia projektu akademickiego.",
    )
    add_heading(doc, "1.1. Cel i zakres projektu", 2)
    add_bullets(
        doc,
        [
            "analiza wybranych właściwości wielowymiarowych szeregów czasowych TEP,",
            "budowa liniowego modelu przestrzeni stanów,",
            "estymacja stanu za pomocą filtra Kalmana,",
            "prognozowanie 41 pomiarów na jeden krok,",
            "porównanie modelu z prognozą naiwną,",
            "wyznaczanie sterowania za pomocą ograniczonego MPC,",
            "zapis historii danych i wyników w PostgreSQL,",
            "komunikacja zdarzeniowa za pomocą MQTT,",
            "demonstracja trybu odtwarzania i pętli zamkniętej.",
        ],
    )
    add_figure(doc, "01_architektura.png", "Rys. 1. Ogólny przepływ danych i wyników w systemie.")

    add_heading(doc, "2. Zbiór danych Tennessee Eastman Process", 1)
    add_body(
        doc,
        "Tennessee Eastman Process jest procesem testowym opisanym przez Downsa i Vogela [1]. Reprezentuje przemysłową instalację chemiczną obejmującą kilka powiązanych aparatów i pętlę recyrkulacji. Zbiór jest często wykorzystywany do badań nad sterowaniem, diagnostyką i analizą danych procesowych, ponieważ zawiera wiele wzajemnie zależnych pomiarów oraz zmienne manipulowane.",
    )
    add_table(
        doc,
        ["Zbiór", "Liczba przebiegów", "Próbki w przebiegu", "Łączna liczba rekordów"],
        [
            ["FaultFree Training", "500", "500", "250 000"],
            ["FaultFree Testing", "500", "960", "480 000"],
        ],
        [2600, 1900, 2100, 2460],
    )
    add_body(
        doc,
        "Każdy rekord zawiera numer przebiegu `simulationRun`, numer próbki `sample`, 41 kolumn `xmeas` oraz 11 kolumn `xmv`. W obecnym etapie użyto danych bez awarii. Pliki zawierające awarie pozostawiono jako możliwość przyszłego rozszerzenia, na przykład o detekcję anomalii.",
    )
    add_heading(doc, "2.1. Charakter sygnałów", 2)
    add_body(
        doc,
        "Poszczególne zmienne różnią się skalą i charakterem zmian. Część sygnałów jest gładka i silnie zależna od poprzedniej wartości, natomiast inne zawierają więcej krótkookresowych wahań. Na rysunku 2 przedstawiono trzy przykładowe pomiary z pierwszych 300 próbek przebiegu testowego. Autokorelacja dla opóźnienia jednego kroku pokazuje, jak silnie bieżąca wartość zależy od poprzedniej.",
    )
    add_figure(doc, "02_przykladowe_szeregi.png", "Rys. 2. Przykładowe szeregi czasowe o różnej skali i autokorelacji.")

    add_heading(doc, "3. Przygotowanie danych i budowa modelu", 1)
    add_heading(doc, "3.1. Standaryzacja", 2)
    add_body(
        doc,
        "Pomiary TEP przyjmują wartości o bardzo różnych rzędach wielkości. Bez skalowania zmienne liczbowe o dużych wartościach dominowałyby obliczenia. Dla każdej kolumny obliczono średnią i odchylenie standardowe na zbiorze treningowym, a następnie zastosowano standaryzację:",
    )
    add_formula(doc, "z = (x - μ) / σ")
    add_body(
        doc,
        "Wartość zero w tej skali oznacza nominalny średni punkt pracy. Parametry skalowania zapisano razem z modelem, dzięki czemu dokładnie te same przekształcenia są używane podczas działania aplikacji. Osobne parametry zastosowano dla pomiarów i wejść sterujących.",
    )
    add_heading(doc, "3.2. Redukcja wymiaru PCA", 2)
    add_body(
        doc,
        "Do utworzenia stanu procesu wykorzystano analizę głównych składowych PCA. Metoda zastępuje 41 skorelowanych pomiarów przez 20 składowych będących ich kombinacjami liniowymi. Wybrane składowe zachowują około 72,4% wariancji zbioru treningowego. Utrata części wariancji jest zamierzona: model staje się prostszy, a część drobnych wahań i szumu nie trafia bezpośrednio do stanu.",
    )
    add_heading(doc, "3.3. Model przestrzeni stanów", 2)
    add_body(
        doc,
        "Dynamikę opisano liniowym modelem dyskretnym. Macierze A i B dopasowano przez regresję liniową z małą regularyzacją. Pary treningowe tworzono tylko wewnątrz jednego `simulationRun`, aby nie wprowadzać sztucznych przejść pomiędzy niezależnymi seriami.",
    )
    add_formula(doc, "x(k+1) = A x(k) + B u(k) + w(k)")
    add_formula(doc, "y(k) = C x(k) + D u(k) + v(k)")
    add_table(
        doc,
        ["Macierz", "Wymiar", "Znaczenie"],
        [
            ["A", "20 x 20", "dynamika stanu"],
            ["B", "20 x 11", "wpływ sterowania na stan"],
            ["C", "41 x 20", "przejście ze stanu do pomiarów"],
            ["D", "41 x 11", "bezpośredni wpływ sterowania; przyjęto zera"],
        ],
        [1500, 1600, 5960],
    )
    add_body(
        doc,
        f"Promień spektralny macierzy A wynosi {summary['spectral_radius']:.3f}. Jest mniejszy od 1, co wskazuje na stabilność dyskretnej dynamiki liniowej. Kowariancje reszt modelu i rekonstrukcji PCA wykorzystano jako macierze szumu procesu oraz pomiaru.",
    )

    add_heading(doc, "4. Estymacja stanu i prognozowanie", 1)
    add_heading(doc, "4.1. Filtr Kalmana", 2)
    add_body(
        doc,
        "Stan PCA nie jest bezpośrednio zapisany w danych, dlatego jest estymowany filtrem Kalmana [2]. Filtr najpierw przewiduje przyszły stan za pomocą modelu, a po otrzymaniu nowego pomiaru koryguje przewidywanie. Wielkość korekty zależy od niepewności modelu i pomiarów.",
    )
    add_formula(doc, "x_pred = A x + B u,     P_pred = A P Aᵀ + Q")
    add_formula(doc, "K = P_pred Cᵀ (C P_pred Cᵀ + R)⁻¹")
    add_formula(doc, "x = x_pred + K (y - C x_pred)")
    add_body(
        doc,
        "Pierwszy stan jest obliczany metodą najmniejszych kwadratów, a kolejne wykorzystują pełną historię filtra. W odróżnieniu od pierwotnej wersji stan nie jest więc szacowany niezależnie od zera dla każdej próbki.",
    )
    add_heading(doc, "4.2. Prognoza na jeden krok", 2)
    add_body(
        doc,
        "Po korekcie stanu przygotowywana jest prognoza wszystkich 41 pomiarów dla próbki `k+1`. Prognoza jest zapisywana przed pojawieniem się wartości rzeczywistej. Gdy następna próbka dotrze do systemu, rekord prognozy zostaje powiązany z rzeczywistym pomiarem i uzupełniony o metryki błędu.",
    )
    add_figure(doc, "03_prognoza_vs_rzeczywistosc.png", "Rys. 3. Porównanie prognoz na jeden krok z wartościami rzeczywistymi.")
    add_body(
        doc,
        "Dla wolniej zmieniających się sygnałów, takich jak `xmeas_7`, model bardzo dobrze śledzi przebieg. Dla sygnałów bardziej zaszumionych prognoza jest gładsza i nie odwzorowuje każdego krótkiego skoku. Jest to typowe dla modelu o zredukowanym wymiarze i filtra, który oddziela dominującą dynamikę od części zakłóceń.",
    )

    add_heading(doc, "5. Ocena jakości prognoz", 1)
    add_heading(doc, "5.1. Zastosowane metryki", 2)
    add_body(doc, "Do oceny zastosowano trzy podstawowe metryki:")
    add_bullets(
        doc,
        [
            "MAE, czyli średni błąd bezwzględny,",
            "RMSE, który mocniej karze duże błędy,",
            "NRMSE, w którym błąd każdej zmiennej jest dzielony przez jej odchylenie standardowe.",
        ],
    )
    add_formula(doc, "MAE = (1/n) Σ |ŷ - y|")
    add_formula(doc, "RMSE = √((1/n) Σ (ŷ - y)²)")
    add_body(
        doc,
        "NRMSE jest najważniejszą metryką zbiorczą, ponieważ 41 pomiarów ma różne jednostki i skale. Zwykły RMSE byłby zdominowany przez zmienne o największych wartościach liczbowych.",
    )
    add_figure(doc, "04_blad_prognozy.png", "Rys. 4. Zmiana oraz rozkład znormalizowanego błędu prognozy.")
    add_heading(doc, "5.2. Porównanie z prognozą naiwną", 2)
    add_body(
        doc,
        "Punktem odniesienia jest prognoza naiwna zakładająca, że następna wartość będzie równa bieżącej. Dla procesów o dużej bezwładności jest to wymagający benchmark, dlatego samo uzyskanie niewielkiego błędu nie wystarcza. Model powinien systematycznie poprawiać wynik naiwny.",
    )
    add_table(
        doc,
        ["Metryka średnia", "Model", "Prognoza naiwna", "Zmiana"],
        [
            ["NRMSE", f"{summary['average_model_nrmse']:.3f}", f"{summary['average_naive_nrmse']:.3f}", f"-{summary['nrmse_improvement_percent']:.1f}%"],
            ["MAE", f"{summary['average_model_mae']:.3f}", f"{summary['average_naive_mae']:.3f}", "wynik niższy o 0,252"],
        ],
        [2500, 1700, 2200, 2660],
    )
    add_figure(doc, "05_model_vs_naiwny.png", "Rys. 5. NRMSE modelu i prognozy naiwnej dla dziesięciu przebiegów testowych.")
    add_body(
        doc,
        "Model osiągnął niższy NRMSE we wszystkich dziesięciu przebiegach. Średnia poprawa wyniosła około 17,4%. Podobny poziom błędu w kolejnych seriach wskazuje, że wynik jest powtarzalny i nie wynika tylko z wyboru jednego korzystnego przebiegu.",
    )

    add_heading(doc, "6. Regulator MPC", 1)
    add_body(
        doc,
        "Model Predictive Control wykorzystuje model procesu do przewidywania przyszłych wyjść i wybiera sekwencję wejść minimalizującą funkcję kosztu [3]. W projekcie zastosowano horyzont ośmiu kroków. Po rozwiązaniu problemu używane jest tylko pierwsze sterowanie, a na kolejnej próbce optymalizacja jest wykonywana ponownie.",
    )
    add_formula(doc, "J = Σ ||y_pred||² + 0,2 Σ ||u||² + 0,5 Σ ||Δu||²")
    add_body(
        doc,
        "Pierwszy składnik dąży do utrzymania standaryzowanych wyjść w pobliżu zera, czyli średniego punktu nominalnego. Drugi ogranicza wielkość sterowania, a trzeci gwałtowne zmiany między krokami. Problem jest rozwiązywany jako liniowe najmniejsze kwadraty z ograniczeniami przy użyciu funkcji `lsq_linear` z biblioteki SciPy [6].",
    )
    add_heading(doc, "6.1. Ograniczenia sterowań", 2)
    add_body(
        doc,
        "Dolne i górne granice wejść wyznaczono z kwantyli 0,5% oraz 99,5% zbioru treningowego. Pozwala to zachować typowe zakresy procesu i jednocześnie ograniczyć wpływ pojedynczych wartości odstających. Wynik optymalizacji jest dodatkowo przycinany do tych granic.",
    )
    add_figure(doc, "06_sterowania_mpc.png", "Rys. 6. Zapisane wejścia TEP, rekomendacje MPC i dopuszczalne zakresy.")
    add_body(
        doc,
        "Rekomendacje MPC nie muszą być identyczne z wejściami zapisanymi w danych. Zbiór TEP powstał przy użyciu własnych regulatorów, natomiast przygotowany MPC działa na uproszczonym modelu i funkcji celu utrzymującej nominalny punkt pracy. Istotne jest, że generowane wartości są stabilne i zawsze pozostają w realistycznych granicach.",
    )

    add_heading(doc, "7. Architektura aplikacji i warstwa danych", 1)
    add_heading(doc, "7.1. PostgreSQL", 2)
    add_body(
        doc,
        "PostgreSQL przechowuje nie tylko surowe pomiary, ale także pełną historię przetwarzania. Schemat jest zorganizowany wokół uruchomienia symulacji. Klucze obce pozwalają jednoznacznie powiązać próbkę, rzeczywiste wejście, przygotowaną prognozę i rekomendację MPC.",
    )
    add_table(
        doc,
        ["Tabela", "Zawartość"],
        [
            ["simulation_runs", "tryb, zbiór, numer przebiegu, status i czas"],
            ["sensor_data", "41 pomiarów dla każdej próbki"],
            ["process_inputs", "11 rzeczywistych lub zastosowanych wejść"],
            ["predictions", "41 prognoz oraz MAE, RMSE i NRMSE"],
            ["input_controls", "11 rekomendacji MPC i status optymalizacji"],
        ],
        [2300, 6760],
    )
    add_heading(doc, "7.2. MQTT", 2)
    add_body(
        doc,
        "MQTT jest lekkim protokołem publish/subscribe [4]. Symulator publikuje zdarzenie na temat `tep/sensors/outputs`, a kontroler publikuje wynik na `tep/control/inputs`. Wiadomość zawiera identyfikator konkretnego rekordu w bazie, dlatego kontroler nie musi wybierać po prostu najnowszej próbki. Pełne dane pozostają w PostgreSQL, natomiast MQTT przenosi małe komunikaty zdarzeniowe.",
    )
    add_heading(doc, "7.3. Tryby działania", 2)
    add_body(doc, "System udostępnia dwa uzupełniające się tryby:")
    add_bullets(
        doc,
        [
            "replay - odtwarza historyczny przebieg, umożliwia uczciwą ocenę prognoz, a sterowanie MPC pozostaje rekomendacją,",
            "closed-loop - wykorzystuje model liniowy jako instalację, odbiera sterowanie MPC przez MQTT i używa go do obliczenia kolejnego stanu.",
        ],
    )

    add_heading(doc, "8. Pętla zamknięta", 1)
    add_body(
        doc,
        "W celu demonstracji rzeczywistego sprzężenia zwrotnego przygotowano tryb, w którym MPC wpływa na symulowany proces. Po każdym pomiarze regulator oblicza nowe wejście, a symulator czeka na wiadomość z odpowiadającym identyfikatorem próbki. Następny stan jest wyznaczany równaniem `Ax + Bu`.",
    )
    add_figure(doc, "07_petla_zamknieta.png", "Rys. 7. Wygaszanie przykładowego zaburzenia z MPC i bez korekty sterowania.")
    add_body(
        doc,
        "W demonstracji MPC szybciej zmniejsza normę wyjść niż wariant bez korekty sterowania. Wynik należy interpretować jako potwierdzenie działania mechanizmu, a nie walidację na fizycznej instalacji. Symulator i regulator używają tego samego modelu, a do pętli nie dodano dodatkowego szumu ani błędu modelowania.",
    )

    add_heading(doc, "9. Implementacja i testy", 1)
    add_heading(doc, "9.1. Wykorzystane technologie", 2)
    add_table(
        doc,
        ["Technologia", "Rola w projekcie"],
        [
            ["Python, NumPy, SciPy", "obliczenia modelu, filtra i optymalizacji"],
            ["Pandas, pyreadr", "wczytanie i przygotowanie danych RData"],
            ["PostgreSQL, SQLAlchemy", "trwały zapis danych i wyników"],
            ["MQTT, Paho, Mosquitto", "komunikacja zdarzeniowa"],
            ["Docker Compose", "uruchomienie bazy i brokera"],
            ["Matplotlib", "powtarzalne wizualizacje"],
        ],
        [2700, 6360],
    )
    add_heading(doc, "9.2. Weryfikacja", 2)
    add_body(
        doc,
        "Testy jednostkowe sprawdzają wymiary i stabilność modelu, poprawność skalowania i odwrotnego skalowania, skończoność wyników filtra Kalmana oraz przestrzeganie granic przez MPC. Dodatkowo wykonano testy integracyjne z rzeczywistym PostgreSQL i Mosquitto uruchomionymi w Dockerze.",
    )
    add_bullets(
        doc,
        [
            "test replay: 12 pomiarów, 12 prognoz, 12 sterowań i 11 ocenionych prognoz,",
            "test closed-loop: 8 pełnych kroków, 8 sterowań, 8 prognoz i 7 ocenionych prognoz,",
            "wszystkie rozwiązania MPC w testach zakończyły się sukcesem,",
            "cztery testy jednostkowe zakończyły się poprawnie.",
        ],
    )

    add_heading(doc, "10. Ograniczenia i możliwe rozszerzenia", 1)
    add_body(
        doc,
        "Zastosowany model jest liniowy, podczas gdy proces TEP ma charakter nieliniowy. Prognoza obejmuje jeden krok, a PCA zachowuje około 72,4% wariancji. Model i ograniczenia są wyznaczone na danych bez awarii. W trybie pętli zamkniętej model instalacji jest taki sam jak model regulatora, dlatego warunki są prostsze niż w rzeczywistym zastosowaniu.",
    )
    add_body(doc, "Naturalne kierunki dalszego rozwoju obejmują:")
    add_bullets(
        doc,
        [
            "analizę zbiorów Faulty i wykrywanie anomalii,",
            "prognozowanie na kilka kroków i osobne metryki dla wybranych sygnałów,",
            "dodanie szumu i błędu modelowania do pętli zamkniętej,",
            "porównanie z modelem autoregresyjnym lub prostą siecią rekurencyjną,",
            "dashboard lub zestaw zapytań analitycznych nad danymi PostgreSQL.",
        ],
    )

    add_heading(doc, "11. Wnioski", 1)
    add_body(
        doc,
        "Zrealizowano kompletny przepływ analizy wielowymiarowego szeregu czasowego: od wczytania danych TEP, przez standaryzację i model przestrzeni stanów, po filtrowanie, prognozowanie, ocenę i zapis wyników. Model poprawił średni NRMSE o około 17,4% względem prognozy naiwnej na dziesięciu niezależnych przebiegach. Oznacza to, że wykorzystanie stanu procesu i wejść sterujących daje dodatkową informację ponad prostym przeniesieniem ostatniej wartości.",
    )
    add_body(
        doc,
        "Regulator MPC generuje wartości w realistycznych granicach, a tryb pętli zamkniętej potwierdza, że decyzja regulatora może zostać przesłana przez MQTT i wykorzystana do obliczenia kolejnego stanu. PostgreSQL przechowuje historię w formie umożliwiającej dalszą analizę i wizualizację. Projekt pozostaje rozwiązaniem akademickim, ale wszystkie jego główne elementy są ze sobą spójne i działają zgodnie z założeniami.",
    )

    add_heading(doc, "Bibliografia", 1)
    references = [
        "[1] J. J. Downs, E. F. Vogel, A plant-wide industrial process control problem, Computers & Chemical Engineering, 17(3), 245-255, 1993, doi: 10.1016/0098-1354(93)80018-I.",
        "[2] R. E. Kalman, A New Approach to Linear Filtering and Prediction Problems, Journal of Basic Engineering, 82(1), 35-45, 1960, doi: 10.1115/1.3662552.",
        "[3] S. J. Qin, T. A. Badgwell, A survey of industrial model predictive control technology, Control Engineering Practice, 11(7), 733-764, 2003, doi: 10.1016/S0967-0661(02)00186-7.",
        "[4] OASIS, MQTT Version 5.0, OASIS Standard, 2019, https://docs.oasis-open.org/mqtt/mqtt/v5.0/mqtt-v5.0.html.",
        "[5] PostgreSQL Global Development Group, PostgreSQL Documentation, https://www.postgresql.org/docs/.",
        "[6] SciPy Developers, scipy.optimize.lsq_linear documentation, https://docs.scipy.org/doc/scipy/reference/generated/scipy.optimize.lsq_linear.html.",
    ]
    for reference in references:
        paragraph = doc.add_paragraph(style="Normal")
        paragraph.paragraph_format.left_indent = Cm(0.7)
        paragraph.paragraph_format.first_line_indent = Cm(-0.7)
        paragraph.add_run(reference)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Saved report: {OUTPUT}")


if __name__ == "__main__":
    build_report()
